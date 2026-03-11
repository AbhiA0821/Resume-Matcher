# parser.py - Resume Parser Module
import pdfplumber
import docx
import re

# ─────────────────────────────────────────────
# COMPREHENSIVE SKILLS LIST — ALL DOMAINS
# ─────────────────────────────────────────────
SKILLS_LIST = [
    # ── Software / Programming ──
    "python", "java", "javascript", "typescript", "c", "c++", "c#",
    "r", "scala", "go", "rust", "kotlin", "swift", "php", "ruby",
    "sql", "mysql", "sqlite", "postgresql", "mongodb", "firebase",
    "html", "css", "react", "nodejs", "node.js", "angular", "vue",
    "django", "flask", "fastapi", "spring boot", "rest api",
    "docker", "kubernetes", "linux", "git", "github", "devops",
    "aws", "azure", "google cloud", "ci/cd",
    "android", "ios", "flutter", "react native",
    "excel", "powerbi", "tableau", "statistics",

    # ── AI / ML / Data Science ──
    "machine learning", "deep learning", "data science", "data analysis",
    "data analytics", "data visualization", "data preprocessing",
    "feature engineering", "model evaluation", "natural language processing",
    "nlp", "computer vision", "neural network", "reinforcement learning",
    "pandas", "numpy", "scikit-learn", "tensorflow", "keras", "pytorch",
    "matplotlib", "seaborn", "plotly", "streamlit",
    "generative ai", "llm", "prompt engineering", "langchain",
    "gan", "transformer", "bert", "gpt", "stable diffusion",
    "mlops", "recommendation system", "collaborative filtering",
    "business intelligence", "big data", "spark", "hadoop",
    "opencv", "yolo", "image processing",
    "jupyter", "jupyter notebook", "google colab",
    "ollama", "hugging face",

    # ── Mechanical Engineering ──
    "autocad", "auto cad", "catia", "solidworks", "solid works",
    "ansys", "creo", "nx", "unigraphics", "inventor",
    "cad", "cam", "cae", "fea", "fem",
    "3d printing", "additive manufacturing", "rapid prototyping",
    "cnc", "cnc machining", "lathe", "milling", "turning",
    "manufacturing", "production engineering", "industrial engineering",
    "quality control", "quality assurance", "six sigma", "lean manufacturing",
    "gd&t", "tolerance", "metrology",
    "thermodynamics", "heat transfer", "fluid mechanics", "fluid dynamics",
    "hvac", "refrigeration", "air conditioning",
    "automobile", "automotive", "ic engine", "internal combustion",
    "robotics", "mechatronics", "automation",
    "product design", "design engineering", "reverse engineering",
    "welding", "fabrication", "sheet metal",
    "tool design", "jig and fixture", "die design",
    "maintenance engineering", "plant maintenance",
    "piping", "piping design", "pipeline",
    "solidedge", "hypermesh", "abaqus", "nastran",
    "staad", "staad pro",

    # ── Electrical Engineering ──
    "electrical", "electrical engineering",
    "plc", "plc programming", "scada", "hmi", "dcs",
    "power systems", "power electronics", "power generation",
    "vlsi", "verilog", "vhdl", "systemverilog",
    "pcb", "pcb design", "altium", "eagle", "kicad",
    "circuit design", "circuit analysis", "analog circuits", "digital circuits",
    "matlab", "simulink",
    "control systems", "pid control", "automation",
    "transformer", "motor", "generator", "drives",
    "solar energy", "renewable energy", "wind energy",
    "battery management", "bms", "energy storage",
    "high voltage", "protection relay",
    "electrical wiring", "single line diagram",
    "load flow analysis", "short circuit analysis",

    # ── Electronics / ECE ──
    "embedded systems", "embedded c", "embedded",
    "microcontroller", "microprocessor",
    "arduino", "raspberry pi", "esp32", "esp8266", "stm32",
    "fpga", "cpld",
    "iot", "internet of things",
    "mqtt", "modbus", "can bus", "spi", "i2c", "uart",
    "signal processing", "dsp", "image processing",
    "rf", "rf design", "antenna design", "wireless",
    "bluetooth", "wifi", "zigbee", "lora",
    "sensor", "actuator", "transducer",
    "oscilloscope", "multimeter", "spectrum analyzer",

    # ── Civil Engineering ──
    "civil engineering", "civil",
    "structural engineering", "structural analysis", "structural design",
    "rcc design", "steel design", "concrete",
    "revit", "revit architecture", "bim",
    "autocad civil", "civil 3d",
    "construction management", "project management",
    "surveying", "total station", "gps survey",
    "gis", "arcgis", "qgis",
    "geotechnical", "soil mechanics", "foundation design",
    "highway engineering", "road design",
    "water resources", "irrigation", "hydrology",
    "environmental engineering", "wastewater treatment",
    "cost estimation", "quantity surveying", "billing",
    "ms project",

    # ── Chemical Engineering ──
    "chemical engineering", "chemical",
    "process engineering", "process design", "process simulation",
    "aspen", "aspen plus", "hysys",
    "heat exchanger", "distillation", "absorption",
    "reaction engineering", "reactor design",
    "mass transfer", "unit operations",
    "refinery", "petrochemical", "polymer",
    "safety engineering", "hazop", "process safety",
    "instrumentation", "process instrumentation",

    # ── Other / Business ──
    "finance", "accounting", "financial analysis",
    "marketing", "digital marketing", "seo", "social media",
    "business analyst", "business analysis",
    "project management", "agile", "scrum", "pmp",
    "hr", "human resources", "recruitment",
    "supply chain", "logistics", "erp", "sap",
    "ms office", "microsoft office", "word", "powerpoint",
]

def extract_text_from_pdf(file_path):
    """Extract text from PDF resume"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                # Also extract text from tables (handles column layouts)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        for cell in row:
                            if cell:
                                text += str(cell) + " "
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from Word resume"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def extract_text(file_path):
    """Auto detect file type and extract text"""
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        return ""

def clean_text(text):
    """Clean extracted text"""
    text = text.lower()
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text

def extract_skills(text):
    """Extract skills from resume text"""
    text_clean = clean_text(text)
    found_skills = []
    for skill in SKILLS_LIST:
        if skill.lower() in text_clean:
            found_skills.append(skill)
    return list(set(found_skills))

def extract_experience_years(text):
    """Extract years of experience from resume"""
    text = clean_text(text)
    patterns = [
        r'(\d+)\+?\s*years?\s*of\s*experience',
        r'experience\s*of\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s*experience'
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return int(match.group(1))
    return 0

def extract_education(text):
    """Extract education level"""
    text = clean_text(text)
    if "phd" in text or "doctorate" in text:
        return 3
    elif "master" in text or "mtech" in text or "mba" in text or "msc" in text or "m.tech" in text:
        return 2
    elif "bachelor" in text or "btech" in text or "b.tech" in text or "be " in text or "b.e" in text or "bsc" in text:
        return 1
    else:
        return 0

def detect_job_role(skills):
    """Detect job role from skills"""
    skill_text = " ".join(skills).lower()
    if any(s in skill_text for s in ["autocad", "solidworks", "catia", "ansys", "cad", "cnc", "manufacturing", "hvac", "automobile", "robotics"]):
        return "mechanical engineering"
    elif any(s in skill_text for s in ["plc", "scada", "vlsi", "verilog", "pcb", "power systems", "electrical", "matlab"]):
        return "electrical engineering"
    elif any(s in skill_text for s in ["embedded", "arduino", "iot", "raspberry pi", "microcontroller", "fpga"]):
        return "embedded / iot"
    elif any(s in skill_text for s in ["civil", "structural", "revit", "staad", "construction", "gis"]):
        return "civil engineering"
    elif any(s in skill_text for s in ["chemical", "process engineering", "piping", "aspen"]):
        return "chemical engineering"
    elif any(s in skill_text for s in ["machine learning", "deep learning", "data science", "nlp", "computer vision"]):
        return "data science"
    elif any(s in skill_text for s in ["generative ai", "llm", "mlops"]):
        return "ai / ml"
    elif any(s in skill_text for s in ["python", "java", "react", "full stack", "devops"]):
        return "software development"
    else:
        return "engineering"

def parse_resume(file_path):
    """Full resume parsing - returns all extracted info"""
    text = extract_text(file_path)
    skills = extract_skills(text)
    experience_years = extract_experience_years(text)
    education = extract_education(text)
    job_role = detect_job_role(skills)

    return {
        "raw_text": text,
        "skills": skills,
        "experience_years": experience_years,
        "education_level": education,
        "skill_count": len(skills),
        "job_role": job_role
    }