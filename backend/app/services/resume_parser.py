import os
import pypdf
import pymupdf

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyMuPDF and pypdf fallback."""
    text = ""
    try:
        doc = pymupdf.open(file_path)
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
    except Exception:
        pass

    if not text.strip():
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception:
            pass
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word DOCX file using python-docx and XML fallback."""
    text = ""
    try:
        import docx
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
    except Exception:
        import zipfile
        import xml.etree.ElementTree as ET
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                text = " ".join([elem.text for elem in tree.iter() if elem.text])
        except Exception:
            pass
    return text.strip()

def extract_resume_text(file_path: str) -> str:
    """Detects file extension and extracts text content."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        return ""
